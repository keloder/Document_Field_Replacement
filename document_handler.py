"""
文档处理器 - Word/WPS 文档的读写与替换操作
基于 python-docx 库，支持 .docx 格式文档的打开、处理和保存
保留原始格式（字体、大小、颜色等）进行文本替换

主要功能：
1. 文档打开与保存
2. 段落和表格中的文本替换
3. 格式保留机制
4. 批量文件处理

技术特性：
- 使用 python-docx 库进行 Word 文档操作
- 支持段落和表格单元格中的文本替换
- 保留原始格式（字体、大小、颜色、加粗、斜体等）
- 异常安全处理机制

注意事项：
- 仅支持 .docx 格式（.doc 格式需要转换为 .docx）
- 替换操作基于文本内容，不修改文档结构
- 格式保留基于第一个 run 的样式

作者：keloder
版本：v0.1
"""

import os
from typing import List, Callable
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P


class DocumentHandler:
    @staticmethod
    def open_document(file_path: str) -> Document:
        """
        打开 Word 文档
        
        参数：
        - file_path: 文档文件路径
        
        返回值：
        - Document 对象
        
        异常：
        - 如果文件不存在或格式不支持，会抛出异常
        """
        return Document(file_path)

    @staticmethod
    def save_document(doc: Document, output_path: str):
        """
        保存 Word 文档
        
        参数：
        - doc: 要保存的 Document 对象
        - output_path: 输出文件路径
        
        功能：
        - 自动创建输出目录（如果不存在）
        - 保存文档到指定路径
        """
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        doc.save(output_path)

    @staticmethod
    def process_document(doc: Document, replacer_func: Callable[[str], str]) -> int:
        """
        处理文档中的文本替换
        
        处理范围：
        - 所有段落（paragraphs）
        - 所有表格中的单元格（tables -> rows -> cells -> paragraphs）
        
        算法步骤：
        1. 遍历所有段落，对每个段落进行替换
        2. 遍历所有表格，对每个单元格中的段落进行替换
        3. 统计实际替换的段落数量
        
        参数：
        - doc: 要处理的 Document 对象
        - replacer_func: 替换函数，接受文本并返回替换后的文本
        
        返回值：
        - 实际发生替换的段落数量
        """
        count = 0

        # 处理文档中的所有段落
        for para in doc.paragraphs:
            if para.text:
                original = para.text
                new_text = replacer_func(original)
                # 如果文本发生变化，执行格式保留替换
                if new_text != original:
                    DocumentHandler._replace_text_preserve_formatting(para, original, new_text)
                    count += 1

        # 处理文档中的所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # 处理单元格中的所有段落
                    for para in cell.paragraphs:
                        if para.text:
                            original = para.text
                            new_text = replacer_func(original)
                            # 如果文本发生变化，执行格式保留替换
                            if new_text != original:
                                DocumentHandler._replace_text_preserve_formatting(para, original, new_text)
                                count += 1

        return count

    @staticmethod
    def _copy_run_format(source_run, target_run):
        """
        复制 run 的格式属性
        
        复制的格式包括：
        - 字体名称（font.name）
        - 字体大小（font.size）
        - 加粗（font.bold）
        - 斜体（font.italic）
        - 下划线（font.underline）
        - 字体颜色（font.color.rgb）
        - 高亮颜色（font.highlight_color）
        
        参数：
        - source_run: 源 run（提供格式）
        - target_run: 目标 run（接收格式）
        
        说明：
        - 使用 try-except 确保格式复制不会因异常中断
        - 部分格式属性可能不存在，需要安全处理
        """
        try:
            target_run.font.name = source_run.font.name
        except Exception:
            pass
        try:
            target_run.font.size = source_run.font.size
        except Exception:
            pass
        try:
            target_run.font.bold = source_run.font.bold
        except Exception:
            pass
        try:
            target_run.font.italic = source_run.font.italic
        except Exception:
            pass
        try:
            target_run.font.underline = source_run.font.underline
        except Exception:
            pass
        try:
            if source_run.font.color and source_run.font.color.rgb:
                target_run.font.color.rgb = source_run.font.color.rgb
        except Exception:
            pass
        try:
            if source_run.font.highlight_color:
                target_run.font.highlight_color = source_run.font.highlight_color
        except Exception:
            pass

    @staticmethod
    def _replace_text_preserve_formatting(para: Paragraph, original: str, new_text: str):
        """
        替换段落中的文本，同时保留原始格式
        
        算法步骤：
        1. 检查段落是否有 run（文本运行）
        2. 如果没有 run，直接添加新文本
        3. 如果有 run，保留第一个 run 并删除其他 run
        4. 将第一个 run 的文本替换为新文本
        5. 复制第一个 run 的格式到自身
        
        参数：
        - para: 要处理的段落
        - original: 原文本（用于验证）
        - new_text: 新文本
        
        说明：
        - 此方法会破坏段落中多个 run 的格式差异
        - 但能确保替换后的文本保持基本格式
        """
        # 如果段落没有 run，直接添加新文本
        if not para.runs:
            para.add_run(new_text)
            return

        # 获取段落完整文本，验证原文本是否存在
        full_text = para.text
        if original not in full_text:
            return

        # 获取第一个 run（用于保留格式）
        first_run = para.runs[0]

        # 删除除第一个 run 外的所有 run
        for run in para.runs[1:]:
            run._element.getparent().remove(run._element)
        
        # 将第一个 run 的文本替换为新文本
        first_run.text = new_text

        # 复制第一个 run 的格式到自身（确保格式一致）
        DocumentHandler._copy_run_format(first_run, first_run)

    @staticmethod
    def batch_process_files(file_paths: List[str], replacer_func: Callable[[str], str], output_dir: str = None) -> dict:
        """
        批量处理多个文档文件
        
        处理流程：
        1. 遍历所有文件路径
        2. 对每个文件执行打开、处理、保存操作
        3. 记录每个文件的处理结果
        4. 返回处理结果字典
        
        参数：
        - file_paths: 文件路径列表
        - replacer_func: 替换函数
        - output_dir: 输出目录（None 表示覆盖原文件）
        
        返回值：
        - 字典格式：{文件路径: 处理结果}
        - 处理结果包含：success, count, output, error 等信息
        """
        results = {}

        for file_path in file_paths:
            try:
                # 打开文档
                doc = DocumentHandler.open_document(file_path)
                # 处理文档
                count = DocumentHandler.process_document(doc, replacer_func)

                # 确定输出路径
                if output_dir:
                    os.makedirs(output_dir, exist_ok=True)
                    output_path = os.path.join(output_dir, os.path.basename(file_path))
                else:
                    output_path = file_path

                # 保存文档
                DocumentHandler.save_document(doc, output_path)

                # 记录成功结果
                results[file_path] = {'success': True, 'count': count, 'output': output_path}
            except Exception as e:
                # 记录失败结果
                results[file_path] = {'success': False, 'error': str(e)}

        return results
